# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\dự án TKCSDL")
OUTPUT_PATH = ROOT / "docs" / "progress_report_tkcsdl_2026-08-15.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0xF2, 0xF4, 0xF7)
PALE_BLUE = RGBColor(0xE8, 0xEE, 0xF5)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1F, 0x5F, 0x3A)
GOLD = RGBColor(0x7A, 0x5A, 0x00)
RED = RGBColor(0x9B, 0x1C, 0x1C)


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DDE3", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_table_layout_fixed(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])


def style_table_text(table, header_fill="F2F4F7", header_font_color=BLACK):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=0, after=2, line=1.10)
                for run in p.runs:
                    set_run_font(run, size=10.5)
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, size=10.5, color=header_font_color, bold=True)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        set_paragraph_spacing(p, before=16, after=8, line=1.10)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        set_paragraph_spacing(p, before=12, after=6, line=1.10)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        set_paragraph_spacing(p, before=8, after=4, line=1.10)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.10)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=11, color=BLACK, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=BLACK)
    return p


def add_note_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_layout_fixed(table)
    set_table_borders(table, color="D7DBE2", size="8")
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p1, before=0, after=4, line=1.10)
    r1 = p1.add_run(title)
    set_run_font(r1, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, before=0, after=0, line=1.10)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph()


def add_metadata_rows(doc, rows):
    for label, value in rows:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.10)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=11, color=BLACK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=11, color=BLACK)


def paragraph_bottom_rule(doc, color="BFC7D3", size="10"):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=12, line=1.0)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def build():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ["Normal", "Title", "Subtitle"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for level, size, color in [
        (1, 16, BLUE),
        (2, 13, BLUE),
        (3, 12, DARK_BLUE),
    ]:
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_p, before=0, after=0, line=1.0)
    hr = header_p.add_run("Báo cáo tiến trình TKCSDL | Film Photography Contest Platform")
    set_run_font(hr, size=9, color=GRAY)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, before=0, after=0, line=1.0)
    fr = footer_p.add_run("Cập nhật đến ngày 15/08/2026")
    set_run_font(fr, size=9, color=GRAY)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    r = title.add_run("BÁO CÁO TIẾN TRÌNH DỰ ÁN TKCSDL")
    set_run_font(r, size=23, color=BLACK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=14, line=1.10)
    srun = subtitle.add_run("AI-powered Film Photography Contest Management Platform")
    set_run_font(srun, size=14, color=GRAY)

    add_metadata_rows(
        doc,
        [
            ("Ngày cập nhật", "Thứ Bảy, 15/08/2026"),
            ("Phạm vi theo dõi", "Requirement Analysis -> System Architecture Design -> Conceptual DB -> Logical DB -> Physical DB -> SQL Server Implementation -> Docker Setup -> Seed Data -> Database Testing -> Validation/Traceability -> Project Freeze"),
            ("Công nghệ chính thức", "Microsoft SQL Server, Docker/Docker Compose, SSMS 22"),
            ("Mục tiêu báo cáo", "Tổng hợp cấu trúc đã làm, tiến độ thực tế, phần còn thiếu và hướng hoàn thiện để đẩy bài lên mức 10/10"),
        ],
    )
    paragraph_bottom_rule(doc)

    add_note_box(
        doc,
        "Đánh giá nhanh hiện trạng",
        "Bộ artifact cốt lõi của project đã được xây dựng khá sâu và đã chạy được trên môi trường thật bằng SQL Server trong Docker. Tuy nhiên, nếu xét theo chuẩn một bài nộp học thuật đạt 10/10, nhóm vẫn nên hoàn thiện thêm phần sơ đồ trực quan, mở rộng độ phủ kiểm thử, chốt lại các proposed business rules với giảng viên hoặc nhóm, và đóng gói bộ báo cáo bảo vệ cuối cùng.",
    )

    add_heading(doc, "1. Cấu trúc đã làm", 1)
    add_body(
        doc,
        "Repository hiện đã được tổ chức thành các nhóm artifact rõ ràng để phục vụ cả việc phát triển kỹ thuật lẫn trình bày học thuật. Cấu trúc này giúp truy vết được từ yêu cầu nghiệp vụ đến thiết kế kiến trúc, mô hình dữ liệu, mã SQL triển khai, kiểm thử và trạng thái freeze của project.",
    )

    repo_table = doc.add_table(rows=1, cols=3)
    set_table_layout_fixed(repo_table)
    set_table_borders(repo_table)
    set_col_widths(repo_table, [1.55, 2.20, 2.75])
    repo_headers = ["Thành phần", "Đường dẫn", "Ý nghĩa"]
    for idx, text in enumerate(repo_headers):
        repo_table.cell(0, idx).text = text
    repo_rows = [
        ("Tài liệu phân tích", "docs/", "Gap analysis, requirement baseline, architecture baseline, thiết kế CSDL 3 cấp, validation và project freeze"),
        ("Mã SQL", "database/", "DDL, constraints, indexes, views, functions, procedures, triggers, seed data"),
        ("Môi trường chạy", "docker-compose.yml + docker/", "Khởi tạo SQL Server bằng Docker và script init database"),
        ("Kiểm thử", "tests/", "Các test SQL cho cả success case và failure case"),
        ("Tài liệu hướng dẫn", "README.md", "Cách dựng môi trường, kết nối bằng SSMS 22, seed và chạy test"),
    ]
    for row in repo_rows:
        cells = repo_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(repo_table)
    set_repeat_header(repo_table.rows[0])

    add_heading(doc, "2. Các bước đã làm được", 1)
    add_body(
        doc,
        "Project đã được triển khai theo trình tự khá đúng chuẩn của một bài System Analysis & Design, thay vì nhảy thẳng vào viết SQL. Điều này giúp các quyết định ở tầng physical database bám được nghiệp vụ và có rationale rõ ràng.",
    )

    step_table = doc.add_table(rows=1, cols=4)
    set_table_layout_fixed(step_table)
    set_table_borders(step_table)
    set_col_widths(step_table, [0.55, 1.90, 1.40, 2.65])
    for i, text in enumerate(["STT", "Phase", "Trạng thái", "Kết quả chính"]):
        step_table.cell(0, i).text = text
    step_rows = [
        ("1", "Audit repository", "Hoàn thành", "Xác định repository ban đầu gần như rỗng, chỉ có file đề bài và file phân công"),
        ("2", "Requirement Baseline", "Hoàn thành", "Đã xây glossary, stakeholder analysis, role baseline, AS-IS/TO-BE, core flows, FR, NFR, BR và use case"),
        ("3", "Architecture Baseline", "Hoàn thành", "Đã mô tả system boundary, principles, module decomposition, data ownership, AI boundary, audit, RBAC"),
        ("4", "Conceptual Design", "Hoàn thành", "Đã xác định business entities, cardinality, optionality và rationale cho các quan hệ chính"),
        ("5", "Logical Design", "Hoàn thành", "Đã chuyển sang relational model, resolve N:M, data dictionary và phân tích 3NF"),
        ("6", "Physical Design", "Hoàn thành", "Đã mapping sang Microsoft SQL Server với naming convention, datatype, constraint, index strategy"),
        ("7", "Implementation", "Hoàn thành", "Đã viết full bộ script SQL theo thứ tự khởi tạo, object logic và seed data"),
        ("8", "Docker + SSMS workflow", "Hoàn thành", "Đã chạy SQL Server trong Docker, cấu hình kết nối SSMS 22 và xác minh môi trường thật"),
        ("9", "Database Testing", "Hoàn thành một phần mạnh", "Đã có và chạy pass các test cốt lõi cho submission, verification, judging, result, archive và audit"),
        ("10", "Traceability + Freeze", "Hoàn thành", "Đã tạo RTM, CRUD/Data Ownership, Rule-to-Enforcement, Decision Log và project freeze"),
    ]
    for row in step_rows:
        cells = step_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(step_table)
    set_repeat_header(step_table.rows[0])

    add_heading(doc, "3. Hoàn thiện yêu cầu dự án đến đâu rồi", 1)
    add_body(
        doc,
        "Nếu chỉ xét bộ artifact cốt lõi theo đúng scope kỹ thuật hiện tại, project đang ở mức hoàn thiện cao và có thể đánh giá khoảng 90 đến 92 phần trăm. Nếu xét theo chuẩn một bài nộp học thuật cần vừa đúng kỹ thuật, vừa đẹp về hình thức, vừa thuận lợi cho việc bảo vệ trước giảng viên, mức sẵn sàng hiện tại hợp lý hơn ở khoảng 82 đến 88 phần trăm.",
    )
    add_body(
        doc,
        "Lý do của chênh lệch này là vì phần lõi phân tích, thiết kế, SQL triển khai và kiểm thử đã rất chắc, nhưng một số deliverable phục vụ việc trình bày và phòng vệ học thuật vẫn nên được formalize thêm.",
    )

    status_table = doc.add_table(rows=1, cols=4)
    set_table_layout_fixed(status_table)
    set_table_borders(status_table)
    set_col_widths(status_table, [1.75, 1.05, 1.95, 1.75])
    for i, text in enumerate(["Hạng mục yêu cầu", "Mức độ", "Hiện có", "Còn cần để 10/10"]):
        status_table.cell(0, i).text = text
    status_rows = [
        ("Requirement Analysis", "100%", "Bộ requirement baseline đã hoàn chỉnh và bám domain", "Soát lại wording cuối cùng để đồng bộ với báo cáo tổng"),
        ("System Architecture Design", "90-95%", "Đã có đặc tả đầy đủ cho boundary, module, ownership và AI boundary", "Nên vẽ chính thức Context Diagram, Logical Architecture, Component Diagram, Deployment Diagram"),
        ("Conceptual / Logical / Physical DB", "95-100%", "Đã có tài liệu đầy đủ và tách tầng đúng abstraction", "Nên xuất ERD chính thức để nộp và bảo vệ trực quan"),
        ("SQL Server Implementation", "100%", "Script chạy được trên SQL Server thật trong Docker", "Có thể bổ sung thêm views hoặc procedures nếu giảng viên yêu cầu, nhưng không bắt buộc"),
        ("Docker + SSMS 22", "100%", "Workflow đã hoạt động và đã kiểm chứng", "Có thể thêm screenshot hoặc demo script để tiện báo cáo"),
        ("Seed Data", "90%", "Đủ dùng cho walkthrough các core flows", "Nên tăng độ đa dạng dataset để trình diễn nhiều case phức tạp hơn"),
        ("Database Testing", "75-85%", "Core tests đã pass trên database thật", "Chưa phủ hết toàn bộ negative cases như invalid FK, invalid status/date/weight, delete retention violations"),
        ("Validation & Traceability", "95%", "Đã có RTM, CRUD, Rule-to-Enforcement và Decision Log", "Nên bổ sung ma trận đối chiếu với diagram sau khi bộ sơ đồ được vẽ xong"),
        ("Project Freeze", "90%", "Freeze về mặt repository và kỹ thuật đã thực hiện", "Cần freeze thêm bộ sơ đồ và báo cáo cuối cùng để hoàn thành gói nộp học thuật"),
    ]
    for row in status_rows:
        cells = status_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(status_table)
    set_repeat_header(status_table.rows[0])

    add_heading(doc, "4. Còn thiếu bước nào", 1)
    add_body(
        doc,
        "Điểm quan trọng nhất cần lưu ý là project không còn thiếu phần lõi về mặt kỹ thuật dữ liệu, nhưng vẫn còn thiếu một số bước formalization để biến bộ artifact hiện tại thành một bài nộp thật sự mạnh trước hội đồng hoặc giảng viên.",
    )

    missing_table = doc.add_table(rows=1, cols=3)
    set_table_layout_fixed(missing_table)
    set_table_borders(missing_table)
    set_col_widths(missing_table, [2.05, 2.45, 2.00])
    for i, text in enumerate(["Phần còn thiếu", "Mức độ ảnh hưởng", "Khuyến nghị"]):
        missing_table.cell(0, i).text = text
    missing_rows = [
        ("Bộ sơ đồ chính thức", "Ảnh hưởng cao vì đây là deliverable học thuật quan trọng", "Vẽ ERD conceptual, logical, physical; System Context; Logical Architecture; Component; Deployment"),
        ("Mở rộng test coverage", "Ảnh hưởng cao đến tính thuyết phục khi phản biện", "Bổ sung test cho invalid FK, invalid score, invalid status, invalid date, delete retention và rule boundary"),
        ("Chốt proposed business rules", "Ảnh hưởng trung bình đến tính chính danh của thiết kế", "Đánh dấu rule nào cần xác nhận với giảng viên hoặc với nhóm trước khi đưa vào final report"),
        ("Gói báo cáo cuối cùng", "Ảnh hưởng cao đến chất lượng bài nộp", "Biên tập thành báo cáo cuối có narrative mạch lạc, hình ảnh, bảng biểu và kết luận"),
        ("Gói trình bày bảo vệ", "Ảnh hưởng trung bình nhưng có lợi lớn khi thuyết trình", "Chuẩn bị slide, demo flow trong SSMS, ảnh chụp test pass và câu trả lời cho decision log"),
    ]
    for row in missing_rows:
        cells = missing_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(missing_table, header_fill="E8EEF5")
    set_repeat_header(missing_table.rows[0])

    add_heading(doc, "5. Yêu cầu nào của dự án chưa hoàn thiện hoàn toàn", 1)
    add_body(
        doc,
        "Về bản chất, hầu hết các yêu cầu cốt lõi trong repository hiện tại đã được hoàn thiện. Tuy nhiên, nếu đối chiếu nghiêm ngặt với mục tiêu đạt điểm tối đa, vẫn còn một số yêu cầu nên xem là chưa đóng hoàn toàn.",
    )
    add_body(doc, "1. Chưa có bộ sơ đồ trực quan chính thức dưới dạng file nộp. Hiện mới dừng ở mức specification đủ để vẽ.")
    add_body(doc, "2. Database testing chưa phủ hết toàn bộ các integrity failure cases theo chuẩn kiểm thử sâu. Core tests đã pass, nhưng chưa có đầy đủ script riêng cho invalid FK, invalid score range, invalid weight, invalid status/date và deletion/retention violations.")
    add_body(doc, "3. Một số business rules vẫn được đánh dấu là Proposed Business Rule hoặc Design Decision, nghĩa là về mặt học thuật nên còn một bước xác nhận chính thức trước khi đóng final report.")
    add_body(doc, "4. Chưa có bản báo cáo cuối kỳ hoàn chỉnh và gói slide/presentation để phục vụ việc nộp hoặc bảo vệ bài.")
    add_body(doc, "5. Chưa có lớp bằng chứng minh họa trực quan như screenshot SSMS, screenshot chạy Docker, ảnh test pass hoặc walkthrough minh họa cho các core flows.")

    add_heading(doc, "6. Hướng phát triển tiếp theo để đạt 10/10", 1)
    add_body(
        doc,
        "Để đẩy project từ trạng thái rất tốt sang trạng thái xuất sắc, nhóm nên chuyển trọng tâm từ việc xây thêm logic mới sang việc formalize, tăng độ phủ minh chứng và tối ưu tính bảo vệ học thuật.",
    )

    roadmap_table = doc.add_table(rows=1, cols=4)
    set_table_layout_fixed(roadmap_table)
    set_table_borders(roadmap_table)
    set_col_widths(roadmap_table, [0.85, 2.25, 1.35, 2.05])
    for i, text in enumerate(["Ưu tiên", "Việc cần làm tiếp", "Mục tiêu", "Kết quả mong đợi"]):
        roadmap_table.cell(0, i).text = text
    roadmap_rows = [
        ("P1", "Vẽ bộ diagram chính thức từ các specification đã có", "Đóng bộ deliverable học thuật", "Project có thể nộp và giải thích trực quan ngay"),
        ("P1", "Bổ sung test negative cases còn thiếu", "Tăng độ chặt chẽ của validation", "Nâng sức thuyết phục khi bị hỏi về integrity và boundary cases"),
        ("P1", "Soát và xác nhận Proposed BR / Design Decision", "Giảm rủi ro tranh luận với giảng viên", "Final report có lập luận rõ và không bị xem là tự áp đặt rule"),
        ("P2", "Biên tập báo cáo cuối cùng theo narrative thống nhất", "Tăng chất lượng trình bày", "Bài nộp đọc mượt, nhất quán giữa requirement, architecture và database"),
        ("P2", "Chuẩn bị bộ demo trong SSMS 22", "Hỗ trợ bảo vệ và phản biện", "Có thể mở ngay container, query dữ liệu, chạy test và giải thích live"),
        ("P3", "Bổ sung phụ lục ảnh chụp môi trường và kết quả test", "Tăng minh chứng", "Giảng viên nhìn thấy evidence nhanh mà không cần tự dựng lại"),
    ]
    for row in roadmap_rows:
        cells = roadmap_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(roadmap_table)
    set_repeat_header(roadmap_table.rows[0])

    add_heading(doc, "7. Kết luận", 1)
    add_body(
        doc,
        "Tính đến ngày 15/08/2026, project đã vượt qua giai đoạn khó nhất: hình thành được một bộ artifact nhất quán từ requirement đến SQL Server implementation và đã chạy được trên môi trường thật. Đây là nền tảng rất mạnh để nhóm tự tin hoàn thiện bài theo chuẩn cao.",
    )
    add_body(
        doc,
        "Nếu nhóm tiếp tục hoàn thiện thêm sơ đồ chính thức, mở rộng test coverage, chốt các proposed business rules và biên tập bộ báo cáo cuối cùng một cách chỉn chu, khả năng đưa bài lên mức 10/10 là hoàn toàn khả thi.",
    )

    add_heading(doc, "8. Ưu tiên hành động trong 7 ngày tới", 1)
    add_body(
        doc,
        "Để tận dụng đà hiện tại và biến repository kỹ thuật thành bộ nộp học thuật hoàn chỉnh, nhóm nên đóng ngay ba đầu việc quan trọng nhất dưới đây.",
    )
    quick_table = doc.add_table(rows=1, cols=3)
    set_table_layout_fixed(quick_table)
    set_table_borders(quick_table)
    set_col_widths(quick_table, [0.75, 2.75, 3.00])
    for i, text in enumerate(["Mốc", "Việc cần khóa", "Đầu ra cụ thể"]):
        quick_table.cell(0, i).text = text
    quick_rows = [
        ("D+2", "Vẽ full bộ diagram", "System Context, Architecture, Component, Deployment, Conceptual ERD, Logical ERD, Physical ERD"),
        ("D+4", "Mở rộng test suite", "Bổ sung test invalid FK, invalid status/date/weight, delete retention và boundary cases"),
        ("D+7", "Đóng gói final submission", "Báo cáo cuối, slide bảo vệ, screenshot minh chứng, checklist câu hỏi phản biện"),
    ]
    for row in quick_rows:
        cells = quick_table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table_text(quick_table, header_fill="F2F4F7")
    set_repeat_header(quick_table.rows[0])

    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    build()
